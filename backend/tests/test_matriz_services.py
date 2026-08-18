import csv
import io
import logging
import struct
import sys
import unittest
import zipfile
from asyncio import run
from pathlib import Path
from unittest.mock import patch

from docx import Document
from starlette.datastructures import UploadFile

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services.csv_service import CSV_COLUMNS, gerar_csv_matriz
from app.services.xlsx_service import XLSX_COLUMNS, gerar_xlsx_matriz
from app.services.document_service import (
    _extrair_texto_doc,
    extrair_texto_documento,
    validar_documento,
)
from app.main import StatusPollingAccessFilter
from app.services.llm_service import (
    LLMConversionError,
    _criar_linhas_de_fallback,
    _criar_prompt,
    _normalizar_ordens_da_resposta,
    _preencher_item_padrao_detectado,
    _remover_linhas_nao_fundamentadas,
    _matriz_de_conteudo_json,
    _solicitar_ollama_em_stream,
    _validar_cobertura_dos_blocos,
    converter_blocos_com_ia,
)
from app.services.matrix_structure_service import consolidar_hierarquia_tarefas
from app.services.normalizer_service import normalizar_linhas
from app.services.parser_rules_service import filtrar_regras_por_bloco, preparar_blocos_para_ia
from app.services.pdf_service import _serializar_tabela, limpar_texto_pdf, separar_blocos
from app.schemas.matriz import MatrizLinha, MatrizOutput
from app.services.processing_status import (
    atualizar_processamento,
    concluir_processamento,
    iniciar_processamento,
    obter_processamento,
)
from app.routers.extracao_pdf import (
    _agrupar_blocos,
    _exemplos_parser_do_lote,
    _mensagem_progresso_ollama,
    _proxima_raiz_hta,
    _processar_arquivos,
)


class MatrizServicesTest(unittest.TestCase):
    def test_docx_extraction_preserves_paragraphs_and_table_rows(self):
        documento = Document()
        documento.add_paragraph("1. OBJETIVO")
        documento.add_paragraph("Descrever a atividade operacional.")
        tabela = documento.add_table(rows=2, cols=2)
        tabela.cell(0, 0).text = "Etapa"
        tabela.cell(0, 1).text = "Ação"
        tabela.cell(1, 0).text = "1"
        tabela.cell(1, 1).text = "Verificar o sistema"
        arquivo = io.BytesIO()
        documento.save(arquivo)

        texto = extrair_texto_documento(arquivo.getvalue(), "procedimento.docx")

        self.assertIn("1. OBJETIVO", texto)
        self.assertIn("Descrever a atividade operacional.", texto)
        self.assertIn("Etapa\tAção", texto)
        self.assertIn("1\tVerificar o sistema", texto)

    def test_binary_doc_extraction_reads_compressed_piece_table(self):
        texto_fonte = "ANEXO A\r1. OBJETIVO\rExecutar a rotina."
        texto_codificado = texto_fonte.encode("cp1252")
        posicao_texto = 0x200
        word_document = bytearray(posicao_texto + len(texto_codificado))
        word_document[posicao_texto:] = texto_codificado
        struct.pack_into("<H", word_document, 0x0A, 0x0200)
        limites = struct.pack("<II", 0, len(texto_fonte))
        posicao_codificada = (posicao_texto * 2) | 0x40000000
        descritor = struct.pack("<HIH", 0, posicao_codificada, 0)
        segmentos = limites + descritor
        clx = b"\x02" + struct.pack("<I", len(segmentos)) + segmentos
        struct.pack_into("<II", word_document, 0x1A2, 0, len(clx))

        class ArquivoOleFalso:
            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            @staticmethod
            def exists(nome):
                return nome in {"WordDocument", "1Table"}

            @staticmethod
            def openstream(nome):
                return io.BytesIO(bytes(word_document) if nome == "WordDocument" else clx)

        with patch("app.services.document_service.olefile.OleFileIO", return_value=ArquivoOleFalso()):
            texto = _extrair_texto_doc(b"simulated content")

        self.assertEqual(texto.splitlines(), ["ANEXO A", "1. OBJETIVO", "Executar a rotina."])

    def test_document_validation_rejects_extension_content_mismatch(self):
        with self.assertRaisesRegex(ValueError, "PDF válido"):
            validar_documento("procedimento.pdf", b"Word content")

    def test_ollama_json_response_is_validated(self):
        matriz = _matriz_de_conteudo_json(
            '```json\n{"linhas":[{"ordemBloco":1,"descricao":"Executar rotina","tipoTarefa":"Execução"}]}\n```'
        )

        self.assertEqual(matriz.linhas[0].descricao, "Executar rotina")

    def test_ollama_legacy_ordem_field_is_mapped_to_ordem_bloco(self):
        matriz = _matriz_de_conteudo_json(
            '{"linhas":[{"ordem":9,"descricao":"Atividade","tipoTarefa":"Execução"}]}'
        )

        self.assertEqual(matriz.linhas[0].ordemBloco, 9)

    def test_processing_continues_after_a_failed_llm_batch(self):
        blocos = [
            {
                "ordem": ordem,
                "texto": f"Bloco {ordem}",
                "categoria": "geral",
                "escopo": "documento_principal",
                "palavras_chave": [],
            }
            for ordem in range(1, 10)
        ]
        arquivo = UploadFile(filename="teste.pdf", file=io.BytesIO(b"%PDF-1.4 teste"))

        with (
            patch("app.routers.extracao_pdf.extrair_texto_documento", return_value="texto"),
            patch("app.routers.extracao_pdf.limpar_texto_pdf", return_value="texto"),
            patch("app.routers.extracao_pdf.separar_blocos", return_value=blocos),
            patch("app.routers.extracao_pdf.buscar_parser_rules", return_value=[]),
            patch(
                "app.routers.extracao_pdf.converter_blocos_com_ia",
                side_effect=[
                    LLMConversionError("resposta inválida"),
                    [{"descricao": "Bloco 9", "tipoTarefa": "Informação"}],
                ],
            ),
        ):
            linhas, debug = run(_processar_arquivos([arquivo], incluir_debug=False))

        self.assertEqual([linha["descricao"] for linha in linhas], ["Bloco 9"])
        self.assertEqual(debug["falhas_lotes"][0]["ordens_blocos"], list(range(1, 9)))

    def test_processing_route_accepts_legacy_doc(self):
        blocos = [
            {
                "ordem": 1,
                "texto": "1. OBJETIVO",
                "categoria": "secao_principal",
                "escopo": "documento_principal",
                "palavras_chave": [],
            }
        ]
        assinatura_doc = bytes.fromhex("D0CF11E0A1B11AE1")
        arquivo = UploadFile(filename="anexo.doc", file=io.BytesIO(assinatura_doc + b"content"))

        with (
            patch("app.routers.extracao_pdf.extrair_texto_documento", return_value="1. OBJETIVO") as extrator,
            patch("app.routers.extracao_pdf.limpar_texto_pdf", return_value="1. OBJETIVO"),
            patch("app.routers.extracao_pdf.separar_blocos", return_value=blocos),
            patch("app.routers.extracao_pdf.buscar_parser_rules", return_value=[]),
            patch(
                "app.routers.extracao_pdf.converter_blocos_com_ia",
                return_value=[
                    {
                        "ordemBloco": 1,
                        "itemPadrao": "1.",
                        "descricao": "OBJETIVO",
                        "tipoTarefa": "Título/Subtítulo",
                        "subtarefaHTA": "",
                        "descricaoTarefa": "",
                    }
                ],
            ),
        ):
            linhas, _debug = run(_processar_arquivos([arquivo], incluir_debug=False))

        extrator.assert_called_once_with(assinatura_doc + b"content", "anexo.doc")
        self.assertEqual(linhas[0]["descricao"], "OBJETIVO")

    def test_prompt_uses_examples_from_selected_parser_rules(self):
        regras = [
            {
                "nome": "Como fazer",
                "exemplo_entrada": "COMO FAZER: Registrar a atividade.",
                "exemplo_saida_json": [{"descricao": "Registrar a atividade.", "tipoTarefa": "Execução"}],
            },
            {"nome": "Sem exemplo"},
        ]

        exemplos = _exemplos_parser_do_lote(regras)
        prompt = _criar_prompt([{"ordem": 9}], regras, exemplos, {"filename": "teste.pdf", "file_order": 1})

        self.assertEqual(len(exemplos), 1)
        self.assertEqual(prompt["exemplos_parser"][0]["entrada"], "COMO FAZER: Registrar a atividade.")
        self.assertNotIn("nome", prompt["orientacoes_parser"][0])
        self.assertNotIn("exemplos_rag", prompt)
        self.assertEqual(prompt["ordensBlocoPermitidas"], [9])

    def test_parser_adds_matching_rule_orientation_to_each_block(self):
        blocos = [
            {"ordem": 1, "texto": "COMO FAZER: Registrar a atividade.", "escopo": "geral", "categoria": "como_fazer"}
        ]
        regras = [
            {
                "nome": "Como fazer",
                "escopo": "geral",
                "categoria": "como_fazer",
                "tipo_tarefa": "Execução",
                "padrao_regex": r"(?i)^COMO FAZER:",
                "ordem": 1,
            }
        ]

        resultado = preparar_blocos_para_ia(blocos, regras)

        self.assertEqual(resultado[0]["orientacao_parser"]["regra"], "Como fazer")
        self.assertEqual(resultado[0]["orientacao_parser"]["tipoTarefa"], "Execução")

    def test_parser_does_not_apply_unrelated_ignore_rule(self):
        blocos = [{"ordem": 1, "texto": "1. OBJETIVO", "escopo": "documento_principal", "categoria": "objetivo"}]
        regras = [
            {
                "nome": "Sumário inicial",
                "escopo": "documento_principal",
                "categoria": "ruido_sumario",
                "tipo_tarefa": "Ignorar",
                "padrao_regex": r"(?i)^\s*1\.\s*OBJETIVO\s*$",
                "ordem": 1,
            }
        ]

        resultado = preparar_blocos_para_ia(blocos, regras)

        self.assertIsNone(resultado[0]["orientacao_parser"]["regra"])

    def test_parser_uses_category_orientation_when_rule_requires_missing_table_columns(self):
        blocos = [{"ordem": 1, "texto": "1- Alinhar a chegada de dutos", "escopo": "tabela_2", "categoria": "atividade_tabela_2"}]
        regras = [
            {
                "nome": "Atividade da Tabela 2",
                "escopo": "tabela_2",
                "categoria": "atividade_tabela_2",
                "tipo_tarefa": "Execução",
                "padrao_regex": r"Técnico de operação",
                "ordem": 1,
            }
        ]

        resultado = preparar_blocos_para_ia(blocos, regras)

        self.assertEqual(resultado[0]["orientacao_parser"]["regra"], "Atividade da Tabela 2")

    def test_ia_coverage_rejects_missing_source_block(self):
        matriz = MatrizOutput(linhas=[MatrizLinha(ordemBloco=1, descricao="Executar", tipoTarefa="Execução")])

        with self.assertRaises(LLMConversionError):
            _validar_cobertura_dos_blocos([{"ordem": 1}, {"ordem": 2}], matriz)

    def test_ia_recovers_only_the_missing_blocks_once(self):
        blocos = [{"ordem": 1, "texto": "Primeiro"}, {"ordem": 2, "texto": "Segundo"}]
        primeira_resposta = MatrizOutput(linhas=[MatrizLinha(ordemBloco=1, descricao="Primeiro", tipoTarefa="Informação")])
        segunda_resposta = MatrizOutput(linhas=[MatrizLinha(ordemBloco=2, descricao="Segundo", tipoTarefa="Informação")])

        with patch(
            "app.services.llm_service._converter_prompt",
            side_effect=[primeira_resposta, segunda_resposta],
        ) as conversor:
            linhas = converter_blocos_com_ia(blocos, [], [], {"filename": "teste.pdf", "file_order": 1})

        self.assertEqual([linha["ordemBloco"] for linha in linhas], [1, 2])
        self.assertEqual(conversor.call_count, 2)
        self.assertEqual([bloco["ordem"] for bloco in conversor.call_args_list[1].args[0]["blocos"]], [2])

    def test_ia_converts_local_batch_positions_to_global_block_orders(self):
        blocos = [{"ordem": 9}, {"ordem": 10}, {"ordem": 11}]
        resposta = MatrizOutput(
            linhas=[
                MatrizLinha(ordemBloco=1, descricao="Primeiro", tipoTarefa="Informação"),
                MatrizLinha(ordemBloco=2, descricao="Segundo", tipoTarefa="Informação"),
                MatrizLinha(ordemBloco=3, descricao="Terceiro", tipoTarefa="Informação"),
            ]
        )

        normalizada = _normalizar_ordens_da_resposta(blocos, resposta)

        self.assertEqual([linha.ordemBloco for linha in normalizada.linhas], [9, 10, 11])

    def test_parser_completes_detected_item_on_section_title(self):
        matriz = MatrizOutput(
            linhas=[MatrizLinha(ordemBloco=1, descricao="3.2 - Atividade", tipoTarefa="Título/Subtítulo")]
        )

        resultado = _preencher_item_padrao_detectado(
            [{"ordem": 1, "itemPadraoDetectado": "3.2"}],
            matriz,
        )

        self.assertEqual(resultado.linhas[0].itemPadrao, "3.2")
        self.assertEqual(resultado.linhas[0].descricao, "Atividade")

    def test_parser_discards_metadata_not_found_in_the_source_block(self):
        matriz = MatrizOutput(
            linhas=[MatrizLinha(ordemBloco=1, descricao="Contrato CSV da Matriz", tipoTarefa="Informação")]
        )

        resultado = _remover_linhas_nao_fundamentadas(
            [{"ordem": 1, "texto": "Atualizar a pressão do sistema."}],
            matriz,
        )

        self.assertEqual(resultado.linhas, [])

    def test_parser_discards_execution_task_invented_by_the_llm(self):
        matriz = MatrizOutput(
            linhas=[
                MatrizLinha(
                    ordemBloco=1,
                    descricao="Abra a tampa do registrador.",
                    tipoTarefa="Execução",
                    descricaoTarefa="Enviar relatório por e-mail.",
                )
            ]
        )

        resultado = _remover_linhas_nao_fundamentadas(
            [{"ordem": 1, "texto": "Abra a tampa do registrador."}],
            matriz,
        )

        self.assertEqual(resultado.linhas, [])

    def test_ia_recovers_missing_global_orders_from_a_local_retry_response(self):
        blocos = [{"ordem": ordem, "texto": f"Bloco {ordem}"} for ordem in range(9, 17)]
        primeira_resposta = MatrizOutput(
            linhas=[
                MatrizLinha(ordemBloco=ordem, descricao=f"Bloco {ordem}", tipoTarefa="Informação")
                for ordem in range(9, 14)
            ]
        )
        resposta_local_com_excesso = MatrizOutput(
            linhas=[
                MatrizLinha(ordemBloco=ordem, descricao=f"Correção {ordem}", tipoTarefa="Informação")
                for ordem in range(1, 9)
            ]
        )

        with patch(
            "app.services.llm_service._converter_prompt",
            side_effect=[primeira_resposta, resposta_local_com_excesso],
        ) as conversor:
            linhas = converter_blocos_com_ia(blocos, [], [], {"filename": "teste.pdf", "file_order": 1})

        self.assertEqual(conversor.call_count, 2)
        self.assertEqual({linha["ordemBloco"] for linha in linhas}, set(range(9, 17)))

    def test_ia_uses_parser_fallback_after_all_coverage_attempts_fail(self):
        bloco = {"ordem": 14, "texto": "4.1.2 - Requisito técnico", "categoria": "secao_principal"}
        resposta_vazia = MatrizOutput(linhas=[])

        with patch(
            "app.services.llm_service._converter_prompt",
            return_value=resposta_vazia,
        ) as conversor:
            linhas = converter_blocos_com_ia([bloco], [], [], {"filename": "teste.pdf", "file_order": 1})

        self.assertEqual(conversor.call_count, 2)
        self.assertEqual(linhas[0]["ordemBloco"], 14)
        self.assertEqual(linhas[0]["tipoTarefa"], "Título/Subtítulo")
        self.assertEqual(linhas[0]["descricao"], "- Requisito técnico")

    def test_ollama_progress_message_covers_retry_state_and_unknown_states(self):
        self.assertIn("omitiu blocos", _mensagem_progresso_ollama("corrigindo_cobertura"))
        self.assertEqual(
            _mensagem_progresso_ollama("estado_novo"),
            "Ollama está atualizando o processamento do lote.",
        )

    def test_ollama_json_without_linhas_has_clear_error(self):
        with self.assertRaises(LLMConversionError) as context:
            _matriz_de_conteudo_json('{"ordem": "1", "texto": "Conteúdo do PDF"}')

        self.assertIn('chave "linhas"', str(context.exception))

    def test_ollama_stream_reports_real_response_activity(self):
        class RespostaStreamFalsa:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc_value, traceback):
                return False

            def raise_for_status(self):
                return None

            def iter_lines(self):
                return iter(
                    [
                        'data: {"choices":[{"delta":{"reasoning":"Analisando o lote"}}]}',
                        'data: {"choices":[{"delta":{"content":"{\\"linhas\\":"}}]}',
                        'data: {"choices":[]}',
                        'data: {"choices":[{"delta":{"content":"[]}"}}]}',
                        "data: [DONE]",
                    ]
                )

        atualizacoes: list[tuple[str, int, int]] = []
        with patch("app.services.llm_service.httpx.stream", return_value=RespostaStreamFalsa()):
            conteudo = _solicitar_ollama_em_stream(
                "http://ollama.test/v1/chat/completions",
                {"model": "modelo-teste"},
                30,
                lambda estado, trechos, caracteres: atualizacoes.append((estado, trechos, caracteres)),
            )

        self.assertEqual(conteudo, '{"linhas":[]}')
        self.assertIn(("conectando", 0, 0), atualizacoes)
        self.assertIn(("aguardando_resposta", 0, 0), atualizacoes)
        self.assertIn(("raciocinando", 0, 0), atualizacoes)
        self.assertIn(("gerando_resposta", 1, 10), atualizacoes)
        self.assertEqual(atualizacoes[-1], ("resposta_completa", 2, 13))

    def test_processing_status_tracks_progress_until_completion(self):
        identificador = "teste-status-123"
        iniciar_processamento(identificador, total_arquivos=2)
        atualizar_processamento(
            identificador,
            "ia",
            "A IA está gerando linhas.",
            arquivo_atual=2,
            total_arquivos=2,
            lote_atual=3,
            total_lotes=4,
            etapas_concluidas=4,
            etapas_totais=10,
        )
        em_andamento = obter_processamento(identificador)
        self.assertIsNotNone(em_andamento)
        self.assertEqual(em_andamento["progresso_percentual"], 40)
        falhas_lotes = [{"lote_order": 2, "ordens_blocos": [9, 10]}]
        concluir_processamento(identificador, "CSV parcial gerado.", falhas_lotes=falhas_lotes)

        processamento = obter_processamento(identificador)

        self.assertIsNotNone(processamento)
        self.assertEqual(processamento["status"], "concluido")
        self.assertEqual(processamento["etapa"], "concluido")
        self.assertEqual(processamento["arquivo_atual"], 2)
        self.assertEqual(processamento["lote_atual"], 3)
        self.assertEqual(processamento["etapas_concluidas"], 10)
        self.assertEqual(processamento["etapas_totais"], 10)
        self.assertEqual(processamento["progresso_percentual"], 100)
        self.assertEqual(processamento["mensagem"], "CSV parcial gerado.")
        self.assertEqual(processamento["falhas_lotes"], falhas_lotes)

    def test_status_polling_access_logs_are_hidden_but_errors_remain_visible(self):
        filtro = StatusPollingAccessFilter()
        sucesso = logging.LogRecord(
            "uvicorn.access",
            logging.INFO,
            __file__,
            0,
            '127.0.0.1 - "GET /api/extracao-pdf/process/teste-status-123/status HTTP/1.1" 200',
            (),
            None,
        )
        erro = logging.LogRecord(
            "uvicorn.access",
            logging.INFO,
            __file__,
            0,
            '127.0.0.1 - "GET /api/extracao-pdf/process/teste-status-123/status HTTP/1.1" 404',
            (),
            None,
        )

        self.assertFalse(filtro.filter(sucesso))
        self.assertTrue(filtro.filter(erro))

    def test_batching_preserves_order_and_respects_limit(self):
        blocos = [
            {"ordem": 1, "texto": "aaaa"},
            {"ordem": 2, "texto": "bbb"},
            {"ordem": 3, "texto": "ccccc"},
        ]

        lotes = _agrupar_blocos(blocos, 7)

        self.assertEqual([[bloco["ordem"] for bloco in lote] for lote in lotes], [[1, 2], [3]])

    def test_pdf_cleaning_and_block_order(self):
        texto = "INTERNA\nPágina 1 de 2\nANEXO B\n1. OBJETIVO\nDefinir o fluxo.\n\nCOMO FAZER\nExecutar a rotina.\nAprovado por: qualidade"

        blocos = separar_blocos(limpar_texto_pdf(texto))

        self.assertEqual([bloco["ordem"] for bloco in blocos], [1, 2, 3, 4])
        self.assertEqual(blocos[1]["categoria"], "secao_principal")
        self.assertEqual(blocos[2]["texto"], "Definir o fluxo.")
        self.assertEqual(blocos[3]["categoria"], "como_fazer")
        self.assertEqual(blocos[3]["escopo"], "anexo_b")
        self.assertNotIn("Aprovado", "\n".join(bloco["texto"] for bloco in blocos))

    def test_pdf_parser_removes_repeated_initial_table_of_contents(self):
        texto = (
            "PE-3UBA-00263\n1. OBJETIVO\n2. APLICAÇÃO\n3. DESCRIÇÃO\n4. REGISTROS\n5. DEFINIÇÕES\n"
            "1. OBJETIVO\nDescrever o processo.\n2. APLICAÇÃO\nAplicar na unidade."
        )

        blocos = separar_blocos(texto)

        self.assertEqual(
            [bloco["texto"].splitlines()[0] for bloco in blocos],
            ["PE-3UBA-00263", "1. OBJETIVO", "Descrever o processo.", "2. APLICAÇÃO", "Aplicar na unidade."],
        )

    def test_pdf_cleaning_removes_lone_underscore_artifact(self):
        self.assertEqual(limpar_texto_pdf("_\nConteudo valido"), "Conteudo valido")

    def test_pdf_parser_does_not_change_scope_for_anexo_mentioned_in_instruction(self):
        texto = 'Tabela 2 - Etapas\n1- Alinhar dutos\nCOMO FAZER: Seguir o Anexo "A" do procedimento.'

        blocos = separar_blocos(texto)

        self.assertEqual(blocos[-1]["escopo"], "tabela_2")

    def test_pdf_parser_splits_list_items_and_resets_table_scope(self):
        texto = (
            "Tabela 2 - Etapas de execução\n"
            "1- Alinhar a chegada de dutos\n"
            "COMO FAZER: Alinhar o sistema.\n"
            "2- Registrar a atividade\n"
            "3.2.3 - Recursos necessários\n"
            "- Lanterna à prova de explosão"
        )

        blocos = separar_blocos(texto)

        self.assertEqual([bloco["escopo"] for bloco in blocos[:4]], ["tabela_2", "tabela_2", "tabela_2", "tabela_2"])
        self.assertEqual(blocos[4]["escopo"], "documento_principal")
        self.assertEqual(blocos[5]["escopo"], "documento_principal")
        self.assertEqual(blocos[1]["categoria"], "atividade_tabela_2")
        self.assertEqual(blocos[1]["contextoTarefa"], {"itemPadrao": "", "subtarefaHTA": "1."})
        self.assertEqual(blocos[4]["contextoTarefa"], {})

    def test_pdf_parser_extracts_hierarchical_item_for_section_titles(self):
        blocos = separar_blocos("3.2 - Atividade\nFluxo das atividades.\n3.2.1 - Responsável\nTécnico de operação.")

        titulos = [bloco for bloco in blocos if bloco["tituloEstrutural"]]
        self.assertEqual([bloco["itemPadraoDetectado"] for bloco in titulos], ["3.2", "3.2.1"])

    def test_pdf_parser_groups_contiguous_informational_list_in_one_block(self):
        blocos = separar_blocos(
            "3.2.3 - Recursos necessários\n"
            "- Lanterna à prova de explosão;\n"
            "- Rádio e telefone;\n"
            "- Veículo."
        )

        self.assertEqual(len(blocos), 2)
        self.assertTrue(blocos[1]["listaAgrupada"])
        self.assertEqual(blocos[1]["categoria"], "lista_informativa")
        self.assertEqual(blocos[1]["texto"].count("\n"), 2)

    def test_pdf_parser_groups_numbered_list_without_turning_items_into_titles(self):
        blocos = separar_blocos(
            "3.2.9.2 - Ordem de fechamento\n"
            "1 - Campo Alfa\n\n"
            "2 - Campo Beta\n\n"
            "3 - Campo Gama"
        )

        self.assertEqual(len(blocos), 2)
        self.assertTrue(blocos[1]["listaAgrupada"])
        self.assertEqual(blocos[1]["categoria"], "lista_informativa")
        self.assertFalse(blocos[1]["tituloEstrutural"])

    def test_annex_numbering_distinguishes_index_entries_from_execution_steps(self):
        blocos = separar_blocos(
            limpar_texto_pdf(
                "ANEXO B\n"
                "1.1.2 – Acessar dados do registrador.\n"
                "1.1.1.1 - Abra a tampa do registrador e aperte ENTER.\n"
                "1.1.1.2 - Aperte USER LIST 1.\n"
                "Propriedade da Petrobras PAGE 1 de NUMPAGES\n"
                "EMED-099 Novo Valor percentual: 81,55000"
            )
        )

        linhas = [
            linha
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]
        linha_indice = next(linha for linha in linhas if linha.itemPadrao == "Anexo B - 1.1.2")
        execucoes = [linha for linha in linhas if linha.tipoTarefa == "Execução"]

        self.assertEqual(linha_indice.tipoTarefa, "Informação")
        self.assertEqual(len(execucoes), 3)
        self.assertTrue(all(linha.itemPadrao.startswith("Anexo B - 1.1.1.") for linha in execucoes))
        self.assertTrue(any(linha.descricaoTarefa.startswith("Abrir a tampa") for linha in execucoes))
        self.assertNotIn("Propriedade da Petrobras", "\n".join(bloco["texto"] for bloco in blocos))
        self.assertNotIn("EMED-099", "\n".join(bloco["texto"] for bloco in blocos))

    def test_annex_preserves_item_punctuation_and_discards_interface_suffix(self):
        blocos = separar_blocos(
            "ANEXO B\n"
            "1.1. – Sistema de medição.\n"
            "ANEXO B\n"
            "1.1.1. - Abra a tampa.\n"
            "PE-3UBA-00263 – Versão 03.00 – Padrão Ativo\n\n"
            "USER\n\nLIST\n\nFLOW COMP\n\n1 2 3\n\nO r i f i c e D i a m e t e r"
        )

        titulo = next(bloco for bloco in blocos if bloco["itemPadraoDetectado"] == "Anexo B - 1.1.")
        sufixo = blocos[-5:]

        self.assertEqual(titulo["itemPadraoDetectado"], "Anexo B - 1.1.")
        self.assertTrue(all(bloco["categoria"] == "fragmento_interface" for bloco in sufixo))

    def test_hierarchy_can_continue_after_roots_from_previous_documents(self):
        blocos = separar_blocos(
            "ANEXO B\n"
            "1.1.1 - Procedimento de configuração.\n"
            "ANEXO B\n"
            "1.1 - Sistema\n"
            "1.1.1 - Configuração\n"
            "1.1.1.1 - Abra a tampa e aperte ENTER.\n"
            "1.1.1.2 - Feche a tampa."
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        consolidadas = consolidar_hierarquia_tarefas(blocos, linhas, raiz_inicial=12)
        execucoes = [linha for linha in consolidadas if linha["tipoTarefa"] == "Execução"]

        self.assertEqual(
            [linha["subtarefaHTA"] for linha in execucoes],
            ["12.1.1.1.", "12.1.1.2.", "12.1.1.3."],
        )
        self.assertEqual(_proxima_raiz_hta([{"subtarefaHTA": "11.3.2."}]), 12)

    def test_annex_rebuilds_word_local_numbering_inside_each_procedure(self):
        blocos = separar_blocos(
            "ANEXO B\n"
            "2.2.2. CONTROLADOR E INDICADOR DE NÍVEL\n"
            "Procedimento.\n"
            "1. Escolha a tela desejada.\n"
            "2. Clique no controlador.\n"
            "ANEXO B\n"
            "3. VERIFICAÇÃO E ALTERAÇÃO DOS SETS DAS XV’s\n"
            "Procedimento.\n"
            "1. Escolha no Supervisório a tela correspondente.\n"
            "2. Em seguida, clique no link da XV."
        )

        itens = {
            bloco["texto"]: bloco["itemPadraoDetectado"]
            for bloco in blocos
            if bloco["itemPadraoDetectado"]
        }

        self.assertEqual(itens["1. Escolha a tela desejada."], "Anexo B - 2.2.2.(1.)")
        self.assertEqual(itens["2. Clique no controlador."], "Anexo B - 2.2.2.(2.)")
        self.assertEqual(
            itens["1. Escolha no Supervisório a tela correspondente."],
            "Anexo B - 3.(1.)",
        )
        self.assertEqual(itens["2. Em seguida, clique no link da XV."], "Anexo B - 3.(2.)")

    def test_annex_alphanumeric_group_keeps_items_and_hta_in_the_same_branch(self):
        blocos = separar_blocos(
            "ANEXO B\n"
            "1.1.1 - Procedimento principal\n"
            "1.1.1.1 - Abra a tampa.\n"
            "5B – Recolocação da placa\n"
            "1. Baixar o Porta-Placas.\n"
            "2. Observar as condições da junta."
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        consolidadas = consolidar_hierarquia_tarefas(blocos, linhas, raiz_inicial=12)
        grupo = [linha for linha in consolidadas if str(linha["itemPadrao"]).startswith("Anexo B - 5B")]

        self.assertEqual([linha["itemPadrao"] for linha in grupo], [
            "Anexo B - 5B",
            "Anexo B - 5B(1.)",
            "Anexo B - 5B(2.)",
        ])
        self.assertEqual([linha["subtarefaHTA"] for linha in grupo], [
            "12.1.2.",
            "12.1.2.1.",
            "12.1.2.2.",
        ])
        self.assertEqual(grupo[0]["descricao"], "Recolocação da placa")

    def test_annex_keeps_controller_as_information_and_does_not_promote_procedure_or_action(self):
        blocos = separar_blocos(
            "ANEXO B\n"
            "2.2.2. CONTROLADOR E INDICADOR DE NÍVEL\n"
            "Procedimento.\n"
            "1. Escolha a tela desejada.\n"
            "2. Clique no controlador, defina o set point.\n"
            "3. Acompanhe a atuação da válvula."
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]
        consolidadas = consolidar_hierarquia_tarefas(blocos, linhas, raiz_inicial=12)

        controlador = next(linha for linha in consolidadas if "CONTROLADOR E INDICADOR" in linha["descricao"])
        procedimento = next(linha for linha in consolidadas if linha["descricao"] == "Procedimento.")
        clique = next(linha for linha in consolidadas if linha["descricao"].startswith("Clique no controlador"))

        self.assertEqual(controlador["tipoTarefa"], "Informação")
        self.assertEqual(controlador["subtarefaHTA"], "")
        self.assertEqual(procedimento["tipoTarefa"], "Informação")
        self.assertEqual(procedimento["subtarefaHTA"], "")
        self.assertEqual(clique["tipoTarefa"], "Execução")
        self.assertEqual(clique["subtarefaHTA"], "12.2.2.")

    def test_parser_marks_only_explicit_operational_paragraphs_for_ai(self):
        blocos = separar_blocos(
            "3.2.4 - Itens críticos\n"
            "Como boa prática, recomenda-se coletar uma amostra. O supervisor deverá avaliar o retorno.\n"
            "3.2.5 - Indicadores\n"
            "Não devem ser registrados desvios durante manutenção programada."
        )

        self.assertEqual(blocos[1]["categoria"], "instrucao_operacional")
        self.assertNotEqual(blocos[3]["categoria"], "instrucao_operacional")

    def test_operational_fallback_expands_actions_and_assigns_numeric_hta(self):
        blocos = separar_blocos(
            "2.3 - Controle operacional\n"
            "Recomenda-se coletar uma amostra. O supervisor deverá avaliar o retorno."
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        consolidadas = consolidar_hierarquia_tarefas(blocos, linhas)
        execucoes = [linha for linha in consolidadas if linha["tipoTarefa"] == "Execução"]

        self.assertEqual(len(execucoes), 2)
        self.assertEqual([linha["subtarefaHTA"] for linha in execucoes], ["1.1.", "1.2."])
        self.assertTrue(all(linha["itemPadrao"] == "2.3" for linha in execucoes))
        self.assertEqual(consolidadas[0]["subtarefaHTA"], "1.")

    def test_coordinated_actions_repeat_the_full_source_description(self):
        bloco = {
            "ordem": 1,
            "texto": 'COMO FAZER: Lendo e registrando no "Anexo C" os valores.',
            "categoria": "como_fazer",
            "contextoTarefa": {"itemPadrao": "2.3.1", "subtarefaHTA": "1."},
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(len(linhas), 2)
        self.assertEqual(linhas[0].descricao, linhas[1].descricao)
        self.assertEqual([linha.subtarefaHTA for linha in linhas], ["1.1.", "1.2."])

    def test_shared_complement_is_preserved_for_each_coordinated_action(self):
        bloco = {
            "ordem": 1,
            "texto": 'COMO FAZER: Lendo e registrando no "Anexo C" os valores.',
            "categoria": "como_fazer",
            "contextoTarefa": {"itemPadrao": "2.3.1", "subtarefaHTA": "1."},
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(len(linhas), 2)
        self.assertIn('"Anexo C" os valores', linhas[0].descricaoTarefa)
        self.assertIn('"Anexo C" os valores', linhas[1].descricaoTarefa)

    def test_operational_actions_preserve_conditions_and_resolve_local_objects(self):
        blocos = separar_blocos(
            "4.2 - Resposta operacional\n"
            "Caso haja suspeita de vazamento, deverá ser solicitada avaliação do gasoduto. "
            "Ocorrendo confirmação de vazamento no gasoduto, este deverá ser bloqueado em suas extremidades."
        )
        bloco_operacional = next(
            bloco for bloco in blocos if bloco["categoria"] == "instrucao_operacional"
        )

        tarefas = [
            linha.descricaoTarefa
            for linha in _criar_linhas_de_fallback(bloco_operacional)
        ]

        self.assertTrue(any("caso haja suspeita" in tarefa.lower() for tarefa in tarefas))
        self.assertTrue(any("Bloquear o gasoduto" in tarefa for tarefa in tarefas))
        self.assertTrue(any("ocorrendo confirmação" in tarefa.lower() for tarefa in tarefas))

    def test_likely_missing_r_in_parar_is_recovered_only_with_coordinated_action(self):
        blocos = separar_blocos(
            "4.3 - Contingência\n"
            "Nos casos de grande vazamento, deve-se para os compressores, alinhar para o flare."
        )
        bloco_operacional = next(
            bloco for bloco in blocos if bloco["categoria"] == "instrucao_operacional"
        )

        tarefas = [
            linha.descricaoTarefa
            for linha in _criar_linhas_de_fallback(bloco_operacional)
        ]

        self.assertTrue(any(tarefa.startswith("Parar os compressores") for tarefa in tarefas))
        self.assertTrue(all("grande vazamento" in tarefa for tarefa in tarefas))

    def test_long_numbered_negative_statement_is_information_not_title(self):
        blocos = separar_blocos(
            "4.4.1. Não devem ser registrados desvios quando o sistema estiver em manutenção "
            "programada e formalmente comunicada às áreas envolvidas."
        )

        linhas = _criar_linhas_de_fallback(blocos[0])

        self.assertEqual(blocos[0]["categoria"], "geral")
        self.assertEqual(linhas[0].tipoTarefa, "Informação")
        self.assertEqual(linhas[0].itemPadrao, "4.4.1.")

    def test_operational_contract_does_not_depend_on_llm_paraphrase(self):
        blocos = separar_blocos(
            "4.5 - Resposta a desvios\n"
            "Caso haja vazamento, deverá ser solicitada inspeção do gasoduto. "
            "Ocorrendo confirmação, o gasoduto deverá ser bloqueado."
        )

        with patch("app.services.llm_service._converter_prompt") as converter_mock:
            linhas = converter_blocos_com_ia(blocos, [], [], {"filename": "teste.pdf"})

        converter_mock.assert_not_called()
        execucoes = [linha for linha in linhas if linha["tipoTarefa"] == "Execução"]
        self.assertEqual(len(execucoes), 2)
        self.assertTrue(all(linha["descricaoTarefa"] for linha in execucoes))

    def test_operational_table_uses_only_target_matrix_content(self):
        class TabelaFalsa:
            def extract(self):
                return [
                    ["O QUE FAZER", "EXECUTANTE", "ONDE REGISTRAR"],
                    ["1- Alinhar dutos", "Técnico de operação", "Boletim"],
                    ["COMO FAZER: Alinhar.\nPORQUE FAZER: Garantir fluxo.", None, None],
                ]

        texto = _serializar_tabela(TabelaFalsa())

        self.assertIn("1- Alinhar dutos", texto)
        self.assertNotIn("Técnico de operação", texto)
        self.assertNotIn("Boletim\n", texto)
        self.assertIn("COMO FAZER: Alinhar.\n\nPORQUE FAZER: Garantir fluxo.", texto)

    def test_pass_description_table_preserves_steps_and_continuations(self):
        class TabelaFalsa:
            def extract(self):
                return [
                    ["PASSO", "DESCRIÇÃO", "RESP."],
                    ["1", "Abrir a válvula e confirmar a pressão", "O4"],
                    [None, "antes de prosseguir.", None],
                    ["2", "Fechar a válvula.", "O4"],
                    [None, "Realizar abertura das válvulas", None],
                    ["3", None, None],
                    [None, "e confirmar o alinhamento.", None],
                ]

        texto = _serializar_tabela(TabelaFalsa())

        self.assertIn("PASSO 1 | Abrir a válvula e confirmar a pressão antes de prosseguir.", texto)
        self.assertIn("PASSO 2 | Fechar a válvula.", texto)
        self.assertIn("PASSO 3 | Realizar abertura das válvulas e confirmar o alinhamento.", texto)
        self.assertNotIn("RESP", texto)
        self.assertNotIn("O4", texto)

    def test_numeric_technical_table_is_not_treated_as_operational_steps(self):
        class TabelaFalsa:
            def extract(self):
                return [["1", "10 bar"], ["2", "20 bar"]]

        texto = _serializar_tabela(TabelaFalsa())

        self.assertNotIn("PASSO", texto)
        self.assertIn("1 | 10 bar", texto)

    def test_operational_step_keeps_coordinated_actions_in_the_same_task(self):
        bloco = {
            "ordem": 1,
            "texto": "PASSO 7 | Alinhar a chegada do riser e abrir a válvula XV-01.",
            "categoria": "passo_tabela_operacional",
            "itemPadraoDetectado": "3.2.2.7.",
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual([linha.tipoTarefa for linha in linhas], ["Execução"])
        self.assertEqual(
            [linha.descricaoTarefa for linha in linhas],
            ["Alinhar a chegada do riser e abrir a válvula XV-01."],
        )

    def test_informational_table_preserves_header_and_rows(self):
        bloco = {
            "ordem": 1,
            "texto": (
                "• Quem | O que\n"
                "• COOP | Garantir o cronograma.\n"
                "• O4 | Executar as atividades.\n"
                "• Informar qualquer anormalidade."
            ),
            "categoria": "tabela_tecnica",
            "listaAgrupada": True,
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(
            [linha.tipoTarefa for linha in linhas],
            ["Título/Subtítulo", "Informação", "Informação"],
        )
        self.assertEqual(linhas[0].descricao, "Quem - O que")
        self.assertIn("Informar qualquer anormalidade", linhas[-1].descricao)

    def test_standalone_annex_keeps_conditional_actions_as_complete_rows(self):
        texto = limpar_texto_pdf(
            "2. Ações imediatas e corretivas\n"
            "\uf0b7 Caso haja vazamento, fechar a válvula e informar ao SUOP; "
            "\uf0b7 Após despressurização, abortar a operação e comunicar ao SUOP."
        )
        blocos = separar_blocos(texto, "Anexo A1 - Processo.pdf")
        linhas = [
            linha
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
            if linha.tipoTarefa == "Execução"
        ]

        self.assertEqual(len(linhas), 2)
        self.assertTrue(all(not linha.itemPadrao for linha in linhas))
        self.assertIn("fechar a válvula e informar ao SUOP", linhas[0].descricaoTarefa)
        self.assertIn("abortar a operação e comunicar ao SUOP", linhas[1].descricaoTarefa)

    def test_standalone_annex_splits_conditional_bullets_extracted_on_one_line(self):
        bloco = {
            "ordem": 1,
            "texto": (
                "• Caso haja vazamento, fechar a válvula e informar ao SUOP; "
                "• Após despressurização, abortar a operação e comunicar ao SUOP."
            ),
            "categoria": "instrucao_operacional",
            "anexoStandalone": True,
            "secaoContextual": "2",
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(len(linhas), 2)
        self.assertEqual(
            [linha.descricao for linha in linhas],
            [
                "Caso haja vazamento, fechar a válvula e informar ao SUOP",
                "Após despressurização, abortar a operação e comunicar ao SUOP.",
            ],
        )
        self.assertTrue(all(linha.tipoTarefa == "Execução" for linha in linhas))
        self.assertTrue(all(not linha.itemPadrao for linha in linhas))

    def test_standalone_annex_distinguishes_critical_check_from_maintenance_note(self):
        texto = limpar_texto_pdf(
            "1.2. Preparação do equipamento\n\n"
            "1.2.1. Itens Críticos para Verificação\n\n"
            "Antes de iniciar a preparação, verificar as válvulas conforme a Tabela 1.\n\n"
            "1.8. Retirada do equipamento\n\n"
            "1.8.1. Itens Críticos para Verificação\n\n"
            "Em caso de manutenção da unidade hidráulica, suspender atividade e informar ao SUOP."
        )
        blocos = separar_blocos(texto, "Anexo A1 - Processo.pdf")
        linhas_por_texto = {
            bloco["texto"]: _criar_linhas_de_fallback(bloco)
            for bloco in blocos
        }

        verificacao = next(
            linhas
            for texto_bloco, linhas in linhas_por_texto.items()
            if texto_bloco.startswith("Antes de iniciar")
        )
        manutencao = next(
            linhas
            for texto_bloco, linhas in linhas_por_texto.items()
            if texto_bloco.startswith("Em caso de manutenção")
        )

        self.assertEqual(verificacao[0].tipoTarefa, "Execução")
        self.assertEqual(verificacao[0].itemPadrao, "1.2.1.")
        self.assertEqual(manutencao[0].tipoTarefa, "Informação")

    def test_operational_photo_list_keeps_title_and_each_photo(self):
        blocos = separar_blocos(
            "3.2.2. Sequência das atividades\n"
            "ETAPA 1 - Operação\n"
            "PASSO 13 | Observação: Realizar registros fotográficos:\n"
            "- Foto da tampa aberta.\n"
            "- Foto do equipamento fechado."
        )
        bloco_fotos = next(bloco for bloco in blocos if "Foto da tampa" in bloco["texto"])
        linhas = _criar_linhas_de_fallback(bloco_fotos)

        self.assertEqual(bloco_fotos["categoria"], "passo_tabela_operacional")
        self.assertEqual(
            [linha.tipoTarefa for linha in linhas],
            ["Título/Subtítulo", "Execução", "Execução"],
        )

    def test_reference_to_table_continues_an_open_sentence(self):
        blocos = separar_blocos(
            "1.2.1. Itens críticos\n"
            "Antes de iniciar, verificar conforme indicado na\n\n"
            "Tabela 1 e ilustrado pela Figura 1."
        )

        texto_completo = next(bloco["texto"] for bloco in blocos if "Antes de iniciar" in bloco["texto"])

        self.assertIn("indicado na Tabela 1", texto_completo)

    def test_standalone_alphanumeric_annex_uses_filename_and_local_hierarchy(self):
        blocos = separar_blocos(
            "PROCESSO DE LANÇAMENTO DE PIG\n"
            "1- Etapas de execução da atividade\n"
            "1.1 - Recursos necessários\n"
            "- Fita zebrada;\n"
            "1.2. Despressurização e Drenagem\n"
            "• a. Isolar a área com fita zebrada;\n"
            "• b. Confirmar o fechamento da válvula.",
            "Anexo B1 - Processo de Lançamento de PIG.pdf",
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        consolidadas = consolidar_hierarquia_tarefas(blocos, linhas, raiz_inicial=6)
        padrao = next(linha for linha in consolidadas if linha["tipoTarefa"] == "Padrão/Anexo")
        execucoes = [linha for linha in consolidadas if linha["tipoTarefa"] == "Execução"]
        titulo = next(linha for linha in consolidadas if linha["itemPadrao"] == "1.2.")

        self.assertEqual(padrao["itemPadrao"], "Anexo B1")
        self.assertEqual(padrao["descricao"], "PROCESSO DE LANÇAMENTO DE PIG")
        self.assertEqual(padrao["subtarefaHTA"], "6.")
        self.assertEqual(titulo["subtarefaHTA"], "6.1.")
        self.assertEqual([linha["itemPadrao"] for linha in execucoes], ["1.2.(a)", "1.2.(b)"])
        self.assertEqual([linha["subtarefaHTA"] for linha in execucoes], ["6.1.1.", "6.1.2."])

    def test_operational_stages_restart_hta_without_changing_source_item(self):
        blocos = separar_blocos(
            "3.2.2. Sequência das atividades\n"
            "ETAPA 1 - Lançamento de PIG\n"
            "Essa etapa consiste na primeira metade.\n"
            "PASSO 1 | Abrir a válvula.\n"
            "ETAPA 2 - Recebimento de PIG\n"
            "Essa etapa consiste na segunda metade.\n"
            "PASSO 1 | Fechar a válvula."
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        consolidadas = consolidar_hierarquia_tarefas(blocos, linhas)
        execucoes = [linha for linha in consolidadas if linha["tipoTarefa"] == "Execução"]
        titulos_com_hta = [
            linha for linha in consolidadas
            if linha["tipoTarefa"] == "Título/Subtítulo" and linha["subtarefaHTA"]
        ]

        self.assertEqual([linha["itemPadrao"] for linha in execucoes], ["3.2.2.1.", "3.2.2.1."])
        self.assertTrue(all("(3.2.2.1.)" in linha["descricaoTarefa"] for linha in execucoes))
        self.assertEqual([linha["subtarefaHTA"] for linha in execucoes], ["1.1.", "2.1."])
        self.assertEqual([linha["subtarefaHTA"] for linha in titulos_com_hta], ["1.", "2."])

    def test_operational_responsibility_cells_are_not_exported(self):
        blocos = separar_blocos(
            "3.2.2. Sequência das atividades\n"
            "ETAPA 1 - Lançamento\n"
            "PASSO 1 | Abrir a válvula.\n\n"
            "RESP.\n\nOP-PRA1\n\nBarco de Apoio\n\nCIOP\n\nFSO-CIMA"
        )

        fragmentos = [bloco for bloco in blocos if bloco["categoria"] == "fragmento_interface"]

        self.assertEqual(len(fragmentos), 5)
        linhas_fragmento = [
            linha
            for bloco in fragmentos
            for linha in _criar_linhas_de_fallback(bloco)
        ]
        self.assertTrue(all(linha.tipoTarefa == "Ignorar" and not linha.descricao for linha in linhas_fragmento))

    def test_standalone_annex_groups_dotted_checklist_and_keeps_remission_as_information(self):
        blocos = separar_blocos(
            "1.3. PREPARAÇÃO\n"
            "1.3.1. PREMISSAS\n"
            "1.3.1.1 - Seguir Etapa 1.6. Pressurização.\n"
            "a. Pressão normalizada..............................................\n"
            "b. Válvula alinhada................................................",
            "Anexo A1 - Processo.pdf",
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        remissao = next(linha for linha in linhas if "Seguir Etapa" in linha["descricao"])
        checklist = next(linha for linha in linhas if "Pressão normalizada" in linha["descricao"])

        self.assertEqual(remissao["tipoTarefa"], "Informação")
        self.assertEqual(checklist["tipoTarefa"], "Informação")
        self.assertIn("Válvula alinhada", checklist["descricao"])

    def test_conditional_annex_instruction_with_communicate_is_execution(self):
        blocos = separar_blocos(
            "1.7. SITUAÇÕES ANORMAIS\n"
            "1.7.1 Caso seja necessário, comunicar ao supervisor.\n"
            "1.7.2 Fechar a válvula.",
            "Anexo B1 - Processo.pdf",
        )
        linhas = [
            linha.model_dump()
            for bloco in blocos
            for linha in _criar_linhas_de_fallback(bloco)
        ]

        execucoes = [linha for linha in linhas if linha["tipoTarefa"] == "Execução"]

        self.assertEqual([linha["itemPadrao"] for linha in execucoes], ["1.7.1", "1.7.2"])

    def test_structural_blocks_bypass_ai_and_keep_title_and_content_separate(self):
        blocos = separar_blocos("1. OBJETIVO\nDescrever o processo.")

        with patch("app.services.llm_service._converter_prompt") as conversor:
            titulo = converter_blocos_com_ia([blocos[0]], [], [], {"filename": "teste.pdf"})

        conversor.assert_not_called()
        self.assertEqual(titulo[0]["descricao"], "OBJETIVO")
        self.assertEqual(titulo[0]["tipoTarefa"], "Título/Subtítulo")

    def test_document_header_populates_item_column_and_allows_empty_description(self):
        blocos = separar_blocos(
            "PE-3UBA-00263 – Versão 03.00 – Padrão Ativo\n\n"
            "RECEBIMENTO DE GÁS NATURAL"
        )

        linhas = converter_blocos_com_ia(blocos, [], [], {"filename": "teste.pdf"})
        normalizadas = normalizar_linhas(linhas)

        self.assertEqual(normalizadas[0]["itemPadrao"], "PE-3UBA-00263-03.00 - RECEBIMENTO DE GÁS NATURAL")
        self.assertEqual(normalizadas[0]["descricao"], "")
        self.assertEqual(normalizadas[0]["tipoTarefa"], "Padrão/Anexo")

    def test_normalizer_and_csv_numbering(self):
        linhas = normalizar_linhas(
            [
                {"itemPadrao": "P1", "descricao": " Referência ", "tipoTarefa": "Padrão/Anexo"},
                {
                    "descricao": "Executar\natividade",
                    "tipoTarefa": "Execução",
                    "subtarefaHTA": "Fazer",
                },
                {"descricao": "Ignorar", "tipoTarefa": "Ignorar"},
                {
                    "descricao": "Executar atividade",
                    "tipoTarefa": "Execução",
                    "subtarefaHTA": "Fazer",
                },
            ]
        )

        csv_text = gerar_csv_matriz(linhas)
        rows = list(csv.DictReader(io.StringIO(csv_text.removeprefix("\ufeff")), delimiter=";"))

        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]["ID da Subtarefa"], "")
        self.assertEqual(rows[1]["ID da Subtarefa"], "1")
        self.assertEqual(rows[1]["Descrição"], "Executar\natividade")
        self.assertEqual(
            CSV_COLUMNS,
            ("Item (Padrão)", "Descrição", "Tipo da Tarefa", "ID da Subtarefa", "Subtarefa (HTA)", "Descrição da tarefa"),
        )

    def test_rules_prioritize_matching_scope_and_category(self):
        regras = [
            {"id": 1, "escopo": "geral", "categoria": "geral", "ordem": 1},
            {"id": 2, "escopo": "anexo_b", "categoria": "anexo", "ordem": 5},
            {"id": 3, "escopo": "anexo_b", "categoria": "geral", "ordem": 2},
        ]

        selecionadas = filtrar_regras_por_bloco(regras, "anexo_b", "anexo")

        self.assertEqual([regra["id"] for regra in selecionadas], [2, 1, 3])

    def test_generic_annex_rule_applies_to_alphanumeric_annex_scope(self):
        regras = [
            {"id": 1, "escopo": "geral", "categoria": "geral", "ordem": 1},
            {"id": 2, "escopo": "anexo", "categoria": "anexo_documento", "ordem": 2},
        ]

        selecionadas = filtrar_regras_por_bloco(
            regras,
            "anexo_arquivo_b1",
            "anexo_documento",
        )

        self.assertEqual([regra["id"] for regra in selecionadas], [2, 1])

    def test_doc_extension_accepts_docx_package_from_reference_files(self):
        documento = Document()
        documento.add_paragraph("1. OBJETIVO")
        documento.add_paragraph("Executar a rotina.")
        arquivo = io.BytesIO()
        documento.save(arquivo)

        texto = extrair_texto_documento(arquivo.getvalue(), "referencia.DOC")

        self.assertIn("Executar a rotina.", texto)

    def test_prefixed_annex_filename_creates_standalone_scope(self):
        blocos = separar_blocos(
            "PROCESSO DE MANOBRA\n"
            "1 - Verificar a condição do equipamento.",
            "7_10 Anexo M5 - Operação e Manutenção.pdf",
        )

        self.assertEqual(blocos[0]["categoria"], "anexo_documento")
        self.assertEqual(blocos[0]["escopo"], "anexo_arquivo_m5")
        self.assertTrue(blocos[0]["anexoStandalone"])

    def test_fluxogram_states_are_filtered_but_actions_are_preserved(self):
        blocos = separar_blocos(
            "SIM\n"
            "Resp:\n"
            "Acionar o alarme geral de emergência.",
            "ANEXO E - Fluxograma - Exportação.pdf",
        )

        self.assertEqual(blocos[0]["categoria"], "anexo_documento")
        self.assertEqual(blocos[1]["categoria"], "fragmento_interface")
        self.assertEqual(blocos[2]["categoria"], "instrucao_operacional")
        self.assertTrue(all(bloco["tipoDocumento"] == "fluxograma" for bloco in blocos))

    def test_fragmento_interface_is_ignored_without_remote_rules(self):
        linhas = _criar_linhas_de_fallback({
            "ordem": 1,
            "texto": "SIM",
            "categoria": "fragmento_interface",
        })

        self.assertEqual(linhas[0].tipoTarefa, "Ignorar")
        self.assertEqual(normalizar_linhas([linha.model_dump() for linha in linhas]), [])

    def test_instruction_items_with_hyphen_are_exported_one_per_line(self):
        bloco = {
            "ordem": 1,
            "texto": (
                "Sempre que houver analises, verificar os valores conforme referencia:\n"
                "- Verificacao da quantidade de liners alinhados.\n"
                "- Verificacao da razao do diferencial.\n"
                "- Realizar operacao de retrolavagem."
            ),
            "categoria": "instrucao_operacional",
            "secaoContextual": "3.3.2.4.",
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(len(linhas), 4)
        self.assertTrue(linhas[1].descricao.startswith("- Verificacao"))
        self.assertTrue(linhas[2].descricao.startswith("- Verificacao"))
        self.assertEqual(linhas[-1].tipoTarefa, "Execu\u00e7\u00e3o")

    def test_inline_alphabetical_items_are_separated_and_keep_local_item(self):
        bloco = {
            "ordem": 1,
            "texto": "a) Isolar a area; b) Confirmar o fechamento da valvula.",
            "categoria": "instrucao_operacional",
            "secaoContextual": "3.2.4.",
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(
            [linha.itemPadrao for linha in linhas],
            ["3.2.4.(a)", "3.2.4.(b)"],
        )
        self.assertEqual(len(linhas), 2)
        self.assertTrue(linhas[0].descricao.startswith("a)"))
        self.assertTrue(linhas[1].descricao.startswith("b)"))

    def test_informational_hyphen_list_is_exported_one_per_line(self):
        bloco = {
            "ordem": 1,
            "texto": "- Lanterna;\n- Radio;\n- Veiculo.",
            "categoria": "lista_informativa",
            "listaAgrupada": True,
        }

        linhas = _criar_linhas_de_fallback(bloco)

        self.assertEqual(len(linhas), 3)
        self.assertEqual(
            [linha.descricao for linha in linhas],
            ["- Lanterna;", "- Radio;", "- Veiculo."],
        )

    def test_xlsx_matrix_uses_reference_columns_and_is_a_valid_package(self):
        conteudo = gerar_xlsx_matriz(
            [
                {
                    "subtarefaHTA": "1.1.",
                    "itemPadrao": "1.1.",
                    "descricao": "Verificar a pressão.",
                    "tipoTarefa": "Execução",
                    "descricaoTarefa": "Verificar a pressão. (1.1.)",
                }
            ]
        )

        with zipfile.ZipFile(io.BytesIO(conteudo)) as pacote:
            self.assertIn("xl/worksheets/sheet1.xml", pacote.namelist())
            self.assertIn("xl/styles.xml", pacote.namelist())
            self.assertIn("Subtarefa (HTA)", pacote.read("xl/worksheets/sheet1.xml").decode("utf-8"))

        self.assertEqual(
            XLSX_COLUMNS,
            (
                "Subtarefa (HTA)",
                "Item (Padrão)",
                "Descrição da tarefa (Padrão)",
                "Tipo da tarefa",
                "ID da Subtarefa",
                "Descrição da tarefa (HTA)",
            ),
        )


if __name__ == "__main__":
    unittest.main()
